"""
CVRadar
Local JD Parser — Track 1, zero LLM
src/jd_parser.py

Parses a DOCX job description into structured JDRequirements using
only regex and heuristics. No API call of any kind is made.

v3.1 fix: keyword extraction from full-sentence JDs.
Many real-world JDs (like SatSure's) use plain paragraphs instead
of bullet points. The old parser returned full sentences as "skills"
which never matched any CV text. This version extracts SHORT skill
keywords from sentences regardless of formatting.
"""

import re
from pathlib import Path
from docx import Document

from src.schemas import JDRequirements


# =====================================================
# SECTION HEADER DETECTORS
# =====================================================

_MUST_HEADERS = re.compile(
    r"(must.have|mandatory|required\s+skill|core\s+skill|"
    r"essential|minimum\s+qualif|technical\s+skill|ideal\s+candidate"
    r"|key\s+skill|primary\s+skill)",
    re.I,
)
_NICE_HEADERS = re.compile(
    r"(nice.to.have|preferred|good\s+to\s+have|"
    r"desirable|additional\s+skill|added\s+advantage|bonus|plus)",
    re.I,
)
_EDU_HEADERS = re.compile(
    r"(education|qualification|degree|academic)",
    re.I,
)
_EXP_HEADERS = re.compile(
    r"(experience|work\s+exp|professional\s+exp|years\s+of)",
    re.I,
)
_DOMAIN_HEADERS = re.compile(
    r"(domain|sector|industry|field\s+of)",
    re.I,
)
_CERT_HEADERS = re.compile(
    r"(certif|license|accreditation)",
    re.I,
)

# Bullet / list item detection
_BULLET = re.compile(r"^[\s]*[-•◦▪▸*→]\s+(.+)$")
_NUMBERED = re.compile(r"^[\s]*\d+[\.\)]\s+(.+)$")

# Experience number extraction
_EXP_RANGE = re.compile(
    r"(\d{1,2})\s*(?:to|-|–)\s*(\d{1,2})\s*(?:\+)?\s*years?",
    re.I,
)
_EXP_SINGLE = re.compile(
    r"(\d{1,2})\s*\+?\s*years?\s+(?:of\s+)?(?:experience|exp)",
    re.I,
)

# =====================================================
# NOISE SENTENCES (skip these entirely)
# =====================================================

_NOISE_PATTERNS = [
    re.compile(p, re.I) for p in [
        r"^why us",
        r"opportunity to work",
        r"flat organizational",
        r"we are looking",
        r"the applicant",
        r"should have excellent communication",
        r"team player",
        r"no hierarchies",
        r"if you are interested",
        r"^\s*about\s+",
        r"satsure is",
        r"freedom to work",
        r"innovative ideas",
        r"driven by cutting",
        r"impact on society",
        r"best-in-class",
        r"leave policy",
        r"work independently",
        r"should be able",
        r"culture of",
        r"fast.paced",
        r"startup",
        r"competitive salary",
    ]
]


def _is_noise(text: str) -> bool:
    for pat in _NOISE_PATTERNS:
        if pat.search(text):
            return True
    return False


# =====================================================
# KNOWN TECHNICAL TERMS (for fallback extraction)
# =====================================================

_KNOWN_TECH = re.compile(
    r"\b(QGIS|ArcGIS|Arc\s*GIS|ArcMap|ArcPro|Erdas|Erdas\s*Imagine|ENVI|"
    r"SNAP|GDAL|GDAL/OGR|Google\s+Earth\s+Engine|GEE|"
    r"Remote\s+Sensing|satellite\s+imager\w*|satellite\s+data|"
    r"image\s+classification|supervised\s+classification|unsupervised\s+classification|"
    r"change\s+detection|LULC|land\s+use|land\s+cover|crop\s+classification|"
    r"crop\s+mapping|vegetation\s+index|NDVI|NDWI|spectral\s+analysis|"
    r"raster\s+data|vector\s+data|spatial\s+analysis|geospatial|GIS|"
    r"Python|pandas|geopandas|numpy|scikit.learn|TensorFlow|PyTorch|"
    r"SQL|MySQL|PostgreSQL|PostGIS|R\b|MATLAB|"
    r"Excel|PowerPoint|Tableau|Power\s*BI|"
    r"deep\s+learning|machine\s+learning|neural\s+network|CNN|"
    r"UAV|drone\s+imager\w*|LiDAR|photogrammetry|"
    r"geo.?statistics|spatial\s+statistics|"
    r"AutoCAD|CAD|"
    r"AWS|Azure|GCP|cloud|Docker|Git|"
    r"JavaScript|TypeScript|HTML|CSS|React|Node|"
    r"automation|scripting|programming)\b",
    re.I,
)


def _extract_keywords_from_sentence(sentence: str) -> list:
    """
    Extract short, matchable skill keywords from a long JD sentence.
    Returns a list of keyword strings, or empty list for noise sentences.
    """
    sentence = sentence.strip()
    if not sentence or len(sentence) < 5:
        return []
    if _is_noise(sentence):
        return []

    keywords = []

    # Strategy 1: "knowledge/experience/understanding of X, Y and Z"
    know_pat = re.compile(
        r"(?:knowledge|experience|understanding|proficien\w*|"
        r"exposure|famili\w+|hands.on)\s+(?:of|in|with)\s+(.+?)(?:\.|$)",
        re.I,
    )
    m = know_pat.search(sentence)
    if m:
        raw = m.group(1)
        # Split on commas, "and", "or"
        parts = re.split(r"[,/]|\band\b|\bor\b", raw, flags=re.I)
        for part in parts:
            part = re.sub(
                r"\s*(and\s+other.*|software|tools?|platform|engine|"
                r"(?:is\s+)?(?:an?\s+)?(?:added\s+)?advantage|"
                r"desired|preferred|required|etc\.?|"
                r"is\s+desired|is\s+preferred|is\s+required)\s*$",
                "",
                part.strip(),
                flags=re.I,
            ).strip()
            # Remove leading junk: "the", "a", "an", "other", "at least one:"
            part = re.sub(r"^(the|a|an|other|specifically|at\s+least\s+one\s*:?)\s+", "", part, flags=re.I).strip()
            # Remove trailing colon-separated qualifier like "core GIS software: QGIS" → "QGIS"
            if ":" in part:
                part = part.split(":")[-1].strip()
            part = part.rstrip(".")
            if 2 < len(part) < 50 and not _is_noise(part):
                keywords.append(part)
        if keywords:
            return keywords

    # Strategy 2: extract known tech terms from the sentence
    tech_matches = _KNOWN_TECH.findall(sentence)
    if tech_matches:
        return list(dict.fromkeys(t.strip() for t in tech_matches))

    # Strategy 3: short sentence (≤8 words) = treat as keyword phrase
    words = sentence.split()
    if len(words) <= 8:
        clean = re.sub(r"^(must\s+have|should\s+have|need\s+to\s+have|"
                       r"good\s+to\s+have|nice\s+to\s+have)\s+", "",
                       sentence, flags=re.I).strip()
        clean = clean.rstrip(".")
        if clean and not _is_noise(clean):
            return [clean]

    return []


# =====================================================
# DOCX TEXT EXTRACTION
# =====================================================

def _extract_text_from_docx(file_obj) -> list:
    """Returns list of paragraph text strings."""
    doc = Document(file_obj)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # Also pull from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return lines


def _is_list_item(line: str):
    m = _BULLET.match(line) or _NUMBERED.match(line)
    return m.group(1).strip() if m else None


def _extract_experience_str(text: str) -> str:
    m = _EXP_RANGE.search(text)
    if m:
        return f"{m.group(1)} to {m.group(2)} years"
    m = _EXP_SINGLE.search(text)
    if m:
        return f"{m.group(1)}+ years"
    return ""


def _extract_education_str(lines: list) -> str:
    degree_keywords = [
        "phd", "doctorate", "master", "m.tech", "m.sc", "msc",
        "mba", "bachelor", "b.tech", "b.sc", "bsc", "be ", "b.e",
        "diploma", "degree",
    ]
    for line in lines:
        ll = line.lower()
        if any(kw in ll for kw in degree_keywords):
            return line.strip()
    return ""


def _clean_item(text: str) -> str:
    text = re.sub(r"\s*\(.*?\)", "", text)
    text = text.strip(" .:;,")
    return text


# =====================================================
# MAIN PARSE FUNCTION
# =====================================================

def parse_jd(file_obj) -> JDRequirements:
    """
    Main entry point. Accepts a file-like object (Streamlit UploadedFile).
    Returns a JDRequirements object with SHORT, matchable skill keywords.
    """
    lines = _extract_text_from_docx(file_obj)
    full_text = "\n".join(lines)

    must_have: list = []
    nice_to_have: list = []
    education: str = ""
    experience: str = ""
    domain: str = ""
    certifications: list = []

    edu_lines = []
    exp_lines = []
    domain_lines = []
    current_section = None

    for line in lines:
        stripped = line.strip()
        stripped_lower = stripped.lower()

        # --- Section header detection ---
        if _MUST_HEADERS.search(stripped) and len(stripped) < 120:
            current_section = "must"
            continue
        elif _NICE_HEADERS.search(stripped) and len(stripped) < 80:
            current_section = "nice"
            continue
        elif _EDU_HEADERS.search(stripped) and len(stripped) < 60:
            current_section = "edu"
            continue
        elif _EXP_HEADERS.search(stripped) and len(stripped) < 60:
            current_section = "exp"
            continue
        elif _DOMAIN_HEADERS.search(stripped) and len(stripped) < 60:
            current_section = "domain"
            continue
        elif _CERT_HEADERS.search(stripped) and len(stripped) < 60:
            current_section = "cert"
            continue

        if current_section is None:
            continue

        bullet_text = _is_list_item(stripped)
        text = bullet_text if bullet_text else stripped

        if current_section in ("must", "nice"):
            # Try bullet/list item first (exact short skill)
            if bullet_text:
                cleaned = _clean_item(bullet_text)
                if cleaned and len(cleaned) > 2 and not _is_noise(cleaned):
                    if current_section == "must":
                        must_have.append(cleaned)
                    else:
                        nice_to_have.append(cleaned)
            else:
                # Full sentence — extract keywords from it
                keywords = _extract_keywords_from_sentence(stripped)
                if current_section == "must":
                    must_have.extend(keywords)
                else:
                    nice_to_have.extend(keywords)

        elif current_section == "edu":
            edu_lines.append(stripped)
        elif current_section == "exp":
            exp_lines.append(stripped)
        elif current_section == "domain":
            domain_lines.append(stripped)
        elif current_section == "cert":
            if bullet_text:
                cleaned = _clean_item(bullet_text)
                if cleaned and len(cleaned) > 2:
                    certifications.append(cleaned)

    # --- Post-process ---
    all_exp_text = " ".join(exp_lines)
    experience = _extract_experience_str(all_exp_text)
    if not experience:
        experience = _extract_experience_str(full_text)

    education = _extract_education_str(edu_lines) or _extract_education_str(lines)

    if domain_lines:
        domain = _clean_item(domain_lines[0])

    # --- Fallback: if must_have still empty, scan full text for known tech terms ---
    if not must_have:
        tech_matches = _KNOWN_TECH.findall(full_text)
        seen = set()
        for t in tech_matches:
            t = t.strip()
            if t.lower() not in seen and len(t) > 2:
                seen.add(t.lower())
                must_have.append(t)

    # Deduplicate preserving order
    def dedup(lst):
        seen = set()
        out = []
        for item in lst:
            key = item.lower().strip()
            if key not in seen and len(key) > 1:
                seen.add(key)
                out.append(item)
        return out

    # Normalise common variants so they match CV text correctly
    _NORMALISATIONS = {
        "arc gis": "ArcGIS",
        "arcmap": "ArcGIS",
        "arcgis pro": "ArcGIS",
        "erdas imagine": "Erdas Imagine",
        "erdas": "Erdas Imagine",
        "google earth engine": "Google Earth Engine",
        "google earth": "Google Earth Engine",
        "gee": "GEE",
        "python programming": "Python",
        "python (arcpy)": "Python",
        "remote sensing concepts": "Remote Sensing",
        "remote sensing is": "Remote Sensing",
    }

    def normalise(lst):
        out = []
        for item in lst:
            lower = item.lower().strip()
            out.append(_NORMALISATIONS.get(lower, item))
        return out

    return JDRequirements(
        must_have_skills=dedup(normalise(must_have))[:30],
        nice_to_have_skills=dedup(normalise(nice_to_have))[:20],
        education=education,
        experience=experience,
        domain=domain,
        certifications=dedup(certifications)[:15],
    )
